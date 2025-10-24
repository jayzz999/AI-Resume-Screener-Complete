"""Resume Parsing Utilities with Multi-Format Support

This module provides comprehensive functions for parsing resumes in various formats
(PDF, DOCX) and extracting key information such as skills, education, experience, etc.
"""

import re
import string
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# Third-party imports
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None


class ResumeParseError(Exception):
    """Custom exception for resume parsing errors"""
    pass


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text content from a PDF file using PyPDF2.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a string
        
    Raises:
        ResumeParseError: If PDF extraction fails
    """
    if PdfReader is None:
        raise ResumeParseError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
    
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        return text
    except Exception as e:
        raise ResumeParseError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        ResumeParseError: If DOCX extraction fails
    """
    if Document is None:
        raise ResumeParseError("python-docx is not installed. Install it with: pip install python-docx")
    
    try:
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        return text
    except Exception as e:
        raise ResumeParseError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_path: str) -> str:
    """
    Extract text from resume file (auto-detects format).
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text as a string
        
    Raises:
        ResumeParseError: If file format is unsupported or extraction fails
    """
    path = Path(file_path)
    file_extension = path.suffix.lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ResumeParseError(f"Unsupported file format: {file_extension}")


def clean_and_normalize_text(text: str) -> str:
    """
    Clean and normalize text by removing emails, phone numbers, URLs, and extra whitespace.
    
    Args:
        text: Raw text to clean
        
    Returns:
        Cleaned and normalized text
    """
    # Remove email addresses
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    text = re.sub(email_pattern, '', text)
    
    # Remove phone numbers (various formats)
    phone_patterns = [
        r'\+?\d{1,4}?[-.\s]?\(?\d{1,3}?\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',
        r'\d{3}[-.]?\d{3}[-.]?\d{4}'
    ]
    for pattern in phone_patterns:
        text = re.sub(pattern, '', text)
    
    # Remove URLs
    url_pattern = r'https?://\S+|www\.\S+'
    text = re.sub(url_pattern, '', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text


def extract_skills_section(text: str, skill_keywords: Optional[List[str]] = None) -> List[str]:
    """
    Extract skills from resume text.
    
    Args:
        text: Resume text
        skill_keywords: Optional list of skill keywords to search for
        
    Returns:
        List of identified skills
    """
    if skill_keywords is None:
        # Default skill keywords (can be expanded)
        skill_keywords = [
            'python', 'java', 'javascript', 'c\\+\\+', 'c#', 'ruby', 'php', 'swift', 'kotlin',
            'sql', 'nosql', 'mongodb', 'postgresql', 'mysql', 'oracle',
            'html', 'css', 'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'ci/cd',
            'machine learning', 'deep learning', 'nlp', 'computer vision', 'tensorflow', 'pytorch',
            'data analysis', 'data science', 'pandas', 'numpy', 'scikit-learn',
            'agile', 'scrum', 'jira', 'rest api', 'graphql', 'microservices'
        ]
    
    text_lower = text.lower()
    found_skills = []
    
    for skill in skill_keywords:
        # Use word boundaries to avoid partial matches
        pattern = r'\b' + skill + r'\b'
        if re.search(pattern, text_lower, re.IGNORECASE):
            # Add the skill in its original case from the keyword list
            found_skills.append(skill.replace('\\\\', ''))
    
    # Remove duplicates while preserving order
    seen = set()
    unique_skills = []
    for skill in found_skills:
        skill_lower = skill.lower()
        if skill_lower not in seen:
            seen.add(skill_lower)
            unique_skills.append(skill)
    
    return unique_skills


def extract_education(text: str) -> List[Dict[str, str]]:
    """
    Extract education information from resume text.
    
    Args:
        text: Resume text
        
    Returns:
        List of education entries with degree and institution information
    """
    education_entries = []
    
    # Common degree patterns
    degree_patterns = [
        r'\b(Bachelor|B\.?S\.?|B\.?A\.?|Master|M\.?S\.?|M\.?A\.?|PhD|Ph\.?D\.?|Doctorate)\s+(of|in|degree)?\s*([A-Za-z\s]+)',
        r'\b(Associate|A\.?S\.?|A\.?A\.?)\s+(of|in|degree)?\s*([A-Za-z\s]+)',
    ]
    
    # University/College patterns
    institution_pattern = r'(University|College|Institute|School)\s+of\s+[A-Za-z\s]+|[A-Za-z\s]+(University|College|Institute)'
    
    # Find degrees
    for pattern in degree_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            degree_info = {
                'degree': match.group(0).strip(),
                'field': match.group(3).strip() if len(match.groups()) >= 3 else ''
            }
            education_entries.append(degree_info)
    
    # Find institutions
    institutions = re.findall(institution_pattern, text, re.IGNORECASE)
    
    # Try to match institutions with degrees
    for i, entry in enumerate(education_entries):
        if i < len(institutions):
            entry['institution'] = institutions[i][0] if isinstance(institutions[i], tuple) else institutions[i]
    
    return education_entries


def extract_years_of_experience(text: str) -> Optional[int]:
    """
    Extract years of experience from resume text.
    
    Args:
        text: Resume text
        
    Returns:
        Number of years of experience, or None if not found
    """
    # Patterns for explicit experience mentions
    experience_patterns = [
        r'(\d+)\+?\s*years?\s+of\s+experience',
        r'(\d+)\+?\s*yrs?\s+experience',
        r'experience\s*:\s*(\d+)\+?\s*years?',
    ]
    
    for pattern in experience_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    
    # Try to calculate from work history dates
    # Look for year ranges (e.g., 2018-2021, 2018-Present)
    date_ranges = re.findall(r'(\d{4})\s*[-–—]\s*(\d{4}|Present|Current)', text, re.IGNORECASE)
    
    if date_ranges:
        total_experience = 0
        current_year = 2025  # Update this dynamically if needed
        
        for start, end in date_ranges:
            start_year = int(start)
            end_year = current_year if end.lower() in ['present', 'current'] else int(end)
            total_experience += max(0, end_year - start_year)
        
        return total_experience if total_experience > 0 else None
    
    return None


def preprocess_text_for_nlp(text: str) -> str:
    """
    Preprocess text for NLP tasks (tokenization, lowercasing, removing punctuation).
    
    Args:
        text: Raw text to preprocess
        
    Returns:
        Preprocessed text ready for NLP analysis
    """
    # Convert to lowercase
    text = text.lower()
    
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    # Optional: Remove numbers (uncomment if needed)
    # text = re.sub(r'\d+', '', text)
    
    return text


def parse_resume(file_path: str, skill_keywords: Optional[List[str]] = None) -> Dict:
    """
    Comprehensive resume parsing function that extracts all relevant information.
    
    Args:
        file_path: Path to the resume file
        skill_keywords: Optional list of skill keywords to search for
        
    Returns:
        Dictionary containing parsed resume information
    """
    try:
        # Extract raw text
        raw_text = extract_text(file_path)
        
        # Clean text
        cleaned_text = clean_and_normalize_text(raw_text)
        
        # Extract information
        skills = extract_skills_section(cleaned_text, skill_keywords)
        education = extract_education(cleaned_text)
        years_experience = extract_years_of_experience(cleaned_text)
        
        # Preprocess for NLP
        nlp_text = preprocess_text_for_nlp(cleaned_text)
        
        return {
            'raw_text': raw_text,
            'cleaned_text': cleaned_text,
            'nlp_text': nlp_text,
            'skills': skills,
            'education': education,
            'years_of_experience': years_experience,
            'file_path': file_path
        }
    except Exception as e:
        raise ResumeParseError(f"Failed to parse resume: {str(e)}")


if __name__ == "__main__":
    # Example usage
    print("Resume Parser Utilities Module")
    print("This module provides functions for parsing resumes in PDF and DOCX formats.")
    print("\nRequired dependencies:")
    print("  - PyPDF2: pip install PyPDF2")
    print("  - python-docx: pip install python-docx")
